import json
import os
import tempfile
from datetime import datetime
from urllib.parse import urlparse

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import requests


class DocumentGenerationModule:
    LABOR_KEYWORDS = ("labor", "welding", "blasting", "grinding", "inspection")
    EQUIPMENT_KEYWORDS = ("scaffolding", "machine", "equipment", "rental", "crane", "compressor")

    def __init__(
        self,
        gemini_model_name="gemini-1.5-flash",
        output_folder="outputs/final_reports"
    ):
        self.gemini_model_name = gemini_model_name
        self.output_folder = output_folder
        self.current_lang = "en"
        os.makedirs(self.output_folder, exist_ok=True)

    def translate(self, text, lang='en'):
        if not text:
            return ""

        import re
        text_str = str(text)

        is_paragraph = len(text_str) > 30 or any(kw in text_str for kw in ["Assessment", "Prep Item", "Inspection", "Locate", "Removal", "Hull"])
        if is_paragraph:
            print(f"[REPORT] Original AI text: {text_str}")

        # 1. Translation of complete text
        translated_text = text_str
        if lang == 'bahasa':
            # Protected glossary placeholders
            protected_list = [
                "E7018 Welding Electrodes",
                "E7018",
                "Cutting Torch",
                "Skilled Welder",
                "Copper Slag",
                "Garnet",
                "High Strength Shipbuilding Steel",
                "Grade A Shipbuilding Steel",
                "High-strength Shipbuilding Steel",
                "Tinggi-Strength Shipbuilding Steel",
                "Ultrasonic Thickness Gauge",
                "Sandblasting Machine",
                "IDR",
                "sq.m",
                "rods",
                "sheets",
                "machine-days",
                "kg",
                "hours"
            ]

            placeholders = {}
            temp_text = text_str
            for i, term in enumerate(protected_list):
                placeholder = f"__PROTECTED_TERM_{i}__"
                if term in temp_text:
                    placeholders[placeholder] = term
                    temp_text = temp_text.replace(term, placeholder)

            translations = {
                # Metadata & Header info
                "Marine Technical Services": "Layanan Teknis Kelautan",
                "Hull inspection and repair estimation support": "Dukungan inspeksi lambung dan estimasi perbaikan",
                "Prepared from digital survey findings": "Disiapkan dari temuan survei digital",
                "Contact: operations@marine-technical.local": "Kontak: operations@marine-technical.local",
                "REPAIR ESTIMATE": "ESTIMASI PERBAIKAN",
                "BATCH REPAIR ESTIMATE": "ESTIMASI PERBAIKAN BATCH",
                "Ref": "Ref",
                "Date": "Tanggal",
                "Vessel Name": "Nama Kapal",
                "IMO Number": "Nomor IMO",
                "Document Type": "Jenis Dokumen",
                "Repair Cost Estimate": "Estimasi Biaya Perbaikan",
                "Scope": "Cakupan",
                "Hull defect repair estimate": "Estimasi perbaikan cacat lambung",
                "Prepared By": "Dibuat Oleh",
                "Automated Inspection Workflow": "Alur Kerja Inspeksi Otomatis",
                "Currency": "Mata Uang",
                "Basis": "Dasar",
                "Inspection findings and repair rules": "Temuan inspeksi dan aturan perbaikan",
                "Validity": "Validitas",
                "Subject to onboard verification": "Tunduk pada verifikasi di atas kapal",
                "Total Defects": "Total Cacat",
                "Estimated Total": "Total Estimasi",
                "Material Cost": "Biaya Material",
                "Labor + Equipment": "Tenaga Kerja + Peralatan",
                "Severity Mix": "Campuran Keparahan",
                "Issue Count": "Jumlah Masalah",
                "Basis of Estimate": "Dasar Estimasi",
                "AI review and operator-approved quantities": "Tinjauan AI dan kuantitas yang disetujui operator",
                "Status": "Status",
                "Budgetary estimate": "Estimasi anggaran",
                
                # Headings
                "WORK SCOPE AND COST BREAKDOWN": "CAKUPAN KERJA DAN RINCIAN BIAYA",
                "No.": "No.",
                "Work Description": "Deskripsi Kerja",
                "Service": "Layanan",
                "Material": "Material",
                "Total": "Total",
                "TOTAL ESTIMATED COST": "TOTAL ESTIMASI BIAYA",
                "DEFECT REGISTER": "DAFTAR CACAT",
                "Defect ID": "ID Cacat",
                "Defect Type": "Jenis Cacat",
                "Location": "Lokasi",
                "Severity": "Tingkat Keparahan",
                "Area": "Area",
                "DEFECT PROOF AND COST JUSTIFICATION": "BUKTI CACAT DAN JUSTIFIKASI BIAYA",
                "Approved Item": "Item yang Disetujui",
                "Qty": "Jumlah",
                "Unit Cost": "Biaya Satuan",
                "Line Total": "Total Baris",
                "Approved Total": "Total Disetujui",
                "Image note": "Catatan gambar",
                "NOTES": "CATATAN",
                
                # Footer notes
                "This document is a clean estimate template generated from inspection findings.": "Dokumen ini adalah templat estimasi bersih yang dihasilkan dari temuan inspeksi.",
                "All quantities and locations should be confirmed onboard before commercial issue.": "Semua kuantitas dan lokasi harus dikonfirmasi di atas kapal sebelum penerbitan komersial.",
                "The format intentionally avoids real client identities, billing data, and company names.": "Format ini sengaja menghindari identitas klien nyata, data penagihan, dan nama perusahaan.",
                "Descriptions are kept brief for operational review and approval workflows.": "Deskripsi dibuat singkat untuk tinjauan operasional dan alur kerja persetujuan.",
                "Line items are based on the approved defect repair scope.": "Item baris didasarkan pada cakupan perbaikan cacat yang disetujui.",
                
                # Severity
                "Low": "Rendah",
                "Medium": "Sedang",
                "High": "Tinggi",
                "low": "rendah",
                "medium": "sedang",
                "high": "tinggi",
                
                # Defect types
                "General Defect": "Cacat Umum",
                "Corrosion": "Korosi",
                "Crack": "Retak",
                "Deformation": "Deformasi",
                "Pitting": "Lubang Korosi",
                "Wastage": "Penipisan Plat",
                "Dent": "Penyok",
                "Welding Defect": "Cacat Pengelasan",
                "Hole": "Lubang",
                "Paint Failure": "Kegagalan Cat",
                "Fouling": "Fouling",
                
                # Parts & Locations
                "Hull": "Lambung",
                "Side Shell": "Pelat Sisi",
                "Bottom Shell": "Pelat Alas",
                "Deck": "Geladak",
                "Bulkhead": "Sekat",
                "Frame": "Gading",
                "Transverse Frame": "Gading Melintang",
                "Longitudinal Frame": "Gading Membujur",
                "General area": "Area umum",
                "General Area": "Area Umum",
                
                # Miscellaneous
                "To be confirmed": "Akan dikonfirmasi",
                "pcs": "buah",
                "Prep Item": "Item Persiapan",
                "Repair Item": "Item Perbaikan",
                "General Defect at General area. Repair procedure to be confirmed.": "Cacat Umum pada Area umum. Prosedur perbaikan akan dikonfirmasi.",
                "Proof image could not be embedded": "Gambar bukti tidak dapat disematkan",
                "at": "pada",
                "Item": "Item",
            }

            sentence_translations = {
                # Corrosion
                "Corrosion is the loss of metal due to electrochemical action, especially in salt water environments.":
                    "Korosi adalah terkikisnya logam akibat aksi elektrokimia, terutama di lingkungan air asin.",
                "Ship hulls are highly susceptible to corrosion (rusting) because seawater creates an ideal environment for metal to oxidise.":
                    "Lambung kapal sangat rentan terhadap korosi (berkarat) karena air laut menciptakan lingkungan yang ideal untuk oksidasi logam.",
                "As corrosion progresses, plating thins and may perforate.":
                    "Seiring berjalannya korosi, pelat menipis dan dapat berlubang.",
                "Classification rules typically require replacement of hull plating when wastage exceeds ~15–25%, since localized corrosion or buckling is not acceptable.":
                    "Aturan klasifikasi biasanya mewajibkan penggantian pelat lambung ketika penipisan melebihi ~15–25%, karena korosi terlokalisasi atau tekuk tidak dapat diterima.",
                "Causes include prolonged seawater exposure, trapped moisture, and damaged coatings.":
                    "Penyebabnya meliputi paparan air laut yang berkepanjangan, kelembapan yang terperangkap, dan lapisan pelindung yang rusak.",
                "Use ultrasonic thickness gauges to measure remaining plate thickness and decide if metal replacement is needed.":
                    "Gunakan alat ukur ketebalan ultrasonik untuk mengukur ketebalan pelat yang tersisa dan menentukan apakah penggantian logam diperlukan.",
                "If plating loss is moderate (e.g. up to ~15%), simple cleaning and recoating may suffice.":
                    "Jika penipisan pelat sedang (misalnya hingga ~15%), pembersihan dan pelapisan ulang sederhana saja sudah cukup.",
                "If wastage is higher or holes have formed, proceed to metal replacement.":
                    "Jika tingkat penipisan lebih tinggi or telah terbentuk lubang, lanjutkan ke penggantian logam.",
                "Remove rust and old coatings by abrasive blasting (e.g. with copper-slag or garnet) or power tools to expose clean metal.":
                    "Bersihkan karat dan lapisan cat lama dengan abrasive blasting (misalnya menggunakan copper slag atau garnet) atau peralatan mekanik untuk mengekspos logam yang bersih.",
                "Cut out corroded areas beyond the structurally sound metal (at least to 3/4 plate thickness).":
                    "Potong area yang terkorosi hingga melampaui logam yang sehat secara struktural (setidaknya hingga 3/4 ketebalan pelat).",
                "Make sure edges are smooth and rounded to avoid new stress concentrators.":
                    "Pastikan tepi potongan halus dan membulat untuk menghindari konsentrasi tegangan baru.",
                "Fabricate a new steel patch or insert plate (Grade A shipbuilding steel) to fit the cut-out area.":
                    "Fabrikasi pelat sisipan (insert plate) atau tambalan baja baru (Grade A shipbuilding steel) agar sesuai dengan area potongan.",
                "Align it within one frame spacing if possible.":
                    "Sejajarkan dalam satu jarak gading jika memungkinkan.",
                "The new plate should match or exceed the original thickness.":
                    "Pelat baru harus sesuai atau melebihi ketebalan aslinya.",
                "Position the new plate and weld it continuously around the edges.":
                    "Posisikan pelat baru dan las secara kontinu di sekeliling tepinya.",
                "Use low-hydrogen electrodes (e.g. E7018) and proper preheat (≈150°F) if required.":
                    "Gunakan elektroda rendah hidrogen (misalnya E7018) dan pemanasan awal (preheat) yang tepat (≈150°F) jika diperlukan.",
                "Fillet weld inside and outside the hull if patch plate is used.":
                    "Lakukan las sudut (fillet weld) di bagian dalam dan luar lambung jika menggunakan pelat tambalan.",
                "Ensure welds have full penetration and fair profiles.":
                    "Pastikan hasil las memiliki penetrasi penuh dan profil yang rata.",
                "Grind and smooth weld seams.":
                    "Gerinda dan haluskan kampuh las.",
                "Inspect the repair (e.g. dye-penetrant or ultrasonic NDT) to confirm no defects.":
                    "Inspeksi hasil perbaikan (misalnya dengan dye-penetrant atau ultrasonic NDT) untuk memastikan tidak ada cacat.",
                "Apply primer and anti-corrosive paint (e.g. zinc-chromate primer followed by epoxy/polyurethane topcoats) to protect the repaired area.":
                    "Aplikasikan cat dasar (primer) dan cat anti-korosi (misalnya zinc-chromate primer diikuti dengan cat akhir epoxy/polyurethane) untuk melindungi area yang diperbaiki.",

                # Corrosion Colons
                "Inspection and gauging: Use ultrasonic thickness gauges to measure remaining plate thickness and decide if metal replacement is needed.":
                    "Inspeksi dan pengukuran: Gunakan alat ukur ketebalan ultrasonik untuk mengukur ketebalan pelat yang tersisa dan menentukan apakah penggantian logam diperlukan.",
                "Surface preparation: Remove rust and old coatings by abrasive blasting (e.g. with copper-slag or garnet) or power tools to expose clean metal.":
                    "Persiapan permukaan: Bersihkan karat dan lapisan cat lama dengan abrasive blasting (misalnya menggunakan copper slag atau garnet) atau peralatan mekanik untuk mengekspos logam yang bersih.",
                "Cutting corroded steel: Cut out corroded areas beyond the structurally sound metal (at least to 3/4 plate thickness).":
                    "Pemotongan baja terkorosi: Potong area yang terkorosi hingga melampaui logam yang sehat secara struktural (setidaknya hingga 3/4 ketebalan pelat).",
                "Fitting new plate: Fabricate a new steel patch or insert plate (Grade A shipbuilding steel) to fit the cut-out area.":
                    "Pemasangan pelat baru: Fabrikasi pelat sisipan (insert plate) atau tambalan baja baru (Grade A shipbuilding steel) agar sesuai dengan area potongan.",
                "Welding: Position the new plate and weld it continuously around the edges.":
                    "Pengelasan: Posisikan pelat baru dan las secara kontinu di sekeliling tepinya.",
                "Finishing: Grind and smooth weld seams.":
                    "Finishing: Gerinda dan haluskan kampuh las.",

                # Crack
                "Cracks are splits or fractures in hull steel caused by stress concentration, fatigue, or impact.":
                    "Retak adalah celahan atau patahan pada baja lambung yang disebabkan oleh konsentrasi tegangan, kelelahan logam (fatigue), atau benturan.",
                "They often originate at welds, sharp corners, or other discontinuities under cyclic loading.":
                    "Keretakan sering kali berawal dari lasan, sudut tajam, atau diskontinuitas lainnya di bawah beban siklik.",
                "If cracks open widely or reach the plating edges, classification guidelines typically call for replacing that section with new plate.":
                    "Jika retakan melebar atau mencapai tepi pelat, pedoman klasifikasi biasanya mensyaratkan penggantian bagian tersebut dengan pelat baru.",
                "Causes include repetitive wave slamming, hull flexing, or initial fabrication flaws.":
                    "Penyebabnya meliputi hantaman gelombang yang berulang, kelenturan lambung, atau cacat fabrikasi awal.",
                "Addressing cracks quickly prevents catastrophic failure.":
                    "Penanganan retakan secara cepat dapat mencegah kegagalan katastrofik.",
                "Identify the full extent of the crack (using magnetic-particle or dye-penetrant if needed).":
                    "Identifikasi batas penuh retakan (menggunakan magnetic particle atau dye-penetrant jika diperlukan).",
                "Drill a hole (diameter ≈ plate thickness) at each crack tip to arrest propagation.":
                    "Bor lubang (diameter ≈ ketebalan pelat) di setiap ujung retakan untuk menghentikan penyebaran retak.",
                "Then gouge or cut a “V” groove along the crack into sound metal, removing all oxidized or fractured material.":
                    "Kemudian lakukan gouging atau potong alur \"V\" di sepanjang retakan hingga mencapai logam yang sehat, bersihkan semua material yang teroksidasi atau retak.",
                "Remove any hindering structure (like ceiling panels or stiffeners) to gain access on both sides of the crack.":
                    "Bongkar struktur yang menghalangi (seperti panel langit-langit atau penegar) untuk mendapatkan akses di kedua sisi retakan.",
                "In way of structural members, burn through adjacent welds (≈6” each side) to relieve stress.":
                    "Pada bagian anggota struktural, bersihkan lasan di sekitarnya (≈6 inci di setiap sisi) untuk meredakan tegangan.",
                "With the crack opened in a “V”, begin welding with low-hydrogen stick electrodes (e.g. E7018).":
                    "Dengan retakan yang telah dibuka membentuk alur \"V\", mulailah mengelas menggunakan elektroda las rendah hidrogen (misalnya E7018).",
                "If the crack root gap is too large, build up layers first or use a removable backing strip.":
                    "Jika celah akar (root gap) retakan terlalu besar, buat lapisan penumpuk terlebih dahulu atau gunakan strip penyangga (backing strip) yang dapat dilepas.",
                "After each pass, peen weld beads as recommended, then inspect by grinding and dye-penetrant to ensure no subsurface flaws remain.":
                    "Setelah setiap lintasan las, lakukan peening pada manik las sesuai rekomendasi, lalu inspeksi dengan gerinda dan dye-penetrant untuk memastikan tidak ada cacat di bawah permukaan yang tersisa.",
                "Remove any backing strip after welding, re-weld the back side if needed, and grind smooth.":
                    "Lepaskan backing strip setelah pengelasan, las kembali sisi sebaliknya jika diperlukan, dan gerinda hingga halus.",
                "The repaired area must be continuous with the original plating.":
                    "Area yang diperbaiki harus menyatu dengan pelat asli.",
                "Restore any structural attachments cut earlier.":
                    "Kembalikan semua sambungan struktural yang dipotong sebelumnya.",
                "Finally, coat the welds with primer and paint to restore corrosion protection.":
                    "Akhirnya, lapisi lasan dengan primer dan cat untuk memulihkan perlindungan korosi.",

                # Crack Colons
                "Locate and prepare crack: Identify the full extent of the crack (using magnetic-particle or dye-penetrant if needed).":
                    "Lokalisasi dan persiapan retak: Identifikasi batas penuh retakan (menggunakan magnetic particle atau dye-penetrant jika diperlukan).",
                "Prepare access: Remove any hindering structure (like ceiling panels or stiffeners) to gain access on both sides of the crack.":
                    "Persiapan akses: Bongkar struktur yang menghalangi (seperti panel langit-langit atau penegar) untuk mendapatkan akses di kedua sisi retakan.",
                "Welding procedure: With the crack opened in a “V”, begin welding with low-hydrogen stick electrodes (e.g. E7018).":
                    "Prosedur pengelasan: Dengan retakan yang telah dibuka membentuk alur \"V\", mulailah mengelas menggunakan elektroda las rendah hidrogen (misalnya E7018).",
                "Finish weld: Remove any backing strip after welding, re-weld the back side if needed, and grind smooth.":
                    "Finishing las: Lepaskan backing strip setelah pengelasan, las kembali sisi sebaliknya jika diperlukan, dan gerinda hingga halus.",

                # Deformation
                "Deformation (buckling or warping) refers to out-of-plane bending of hull plating or stiffeners, often from groundings, heavy sea loads, or uneven welding.":
                    "Deformasi (tekuk atau melenting) mengacu pada pembengkokan pelat lambung atau penegar di luar bidangnya, sering kali akibat kandas, beban laut yang berat, atau pengelasan yang tidak merata.",
                "Even moderate buckles can severely impair hull girder strength.":
                    "Bahkan tekuk yang sedang dapat sangat melemahkan kekuatan gelagar lambung.",
                "Classification practice is to correct buckling by cutting out and replacing the distorted metal.":
                    "Praktik klasifikasi adalah memperbaiki tekuk dengan memotong dan mengganti logam yang terdistorsi.",
                "Causes include grounding impacts or welding heat stress.":
                    "Penyebabnya meliputi dampak kandas atau stres panas pengelasan.",
                "Localized set-in plating is sometimes tolerated if small, but significant buckles must be renewed.":
                    "Pelat yang melesak ke dalam secara terlokalisasi terkadang ditoleransi jika kecil, tetapi tekukan yang signifikan harus diganti.",
                "Inspect the deformed region to determine extent.":
                    "Periksa wilayah yang terdeformasi untuk menentukan tingkat kerusakannya.",
                "Mark the buckled area and check if it is localized or spans several frames.":
                    "Tandai area yang menekuk dan periksa apakah kerusakan tersebut terlokalisasi atau mencakup beberapa gading.",
                "Measure any out-of-flatness.":
                    "Ukur setiap penyimpangan kerataan.",
                "Cut away all distorted plate and any buckled girders or stiffeners to sound metal, leaving a margin of good plate (≥¾ thickness).":
                    "Potong semua pelat yang terdistorsi serta gading atau penegar yang menekuk hingga mencapai logam yang sehat, sisakan batas pelat yang baik (≥¾ ketebalan).",
                "Straighten or remove any bent framing if required.":
                    "Luruskan atau lepaskan rangka yang bengkok jika diperlukan.",
                "Fabricate a new steel plate (same or greater thickness) shaped to fit the cut-out.":
                    "Fabrikasi pelat baja baru (ketebalan yang sama atau lebih besar) yang dibentuk agar sesuai dengan potongan.",
                "Position the plate against remaining hull plating.":
                    "Posisikan pelat terhadap pelat lambung yang tersisa.",
                "Tack-weld to hold, then complete weld all around, ensuring good fusion and appropriate preheat.":
                    "Lakukan las ikat (tack-weld) untuk menahan posisinya, kemudian selesaikan pengelasan di sekelilingnya, pastikan fusi yang baik dan pemanasan awal (preheat) yang sesuai.",
                "Use full penetration welds where possible.":
                    "Gunakan las penetrasi penuh jika memungkinkan.",
                "Apply continuing weld passes using E7018 electrodes.":
                    "Lakukan lintasan las berlanjut menggunakan elektroda E7018.",
                "If welding induces distortion, use clamps or minimal heat input techniques.":
                    "Jika pengelasan menyebabkan distorsi, gunakan klem atau teknik masukan panas minimal.",
                "Grind weld seams fair with the hull.":
                    "Gerinda kampuh las agar rata dengan lambung.",
                "In cases of extensive buckling, additional strengthening (doublers or stiffeners) may be applied to reduce stress.":
                    "Untuk tekuk yang luas, penguatan tambahan (doubler plate atau penegar) dapat diterapkan untuk mengurangi tegangan.",
                "Inspect the repair by visual and NDT methods.":
                    "Inspeksi hasil perbaikan secara visual dan dengan metode NDT.",
                "Apply primer and finish coats to the new steel to restore corrosion protection.":
                    "Aplikasikan cat dasar (primer) dan cat akhir pada baja baru untuk memulihkan perlindungan korosi.",
                "Ensure the deck or hull surface is faired back to its original contour.":
                    "Pastikan permukaan geladak atau lambung diratakan kembali ke kontur aslinya.",

                # Deformation Colons
                "Assessment: Inspect the deformed region to determine extent.":
                    "Penilaian: Periksa wilayah yang terdeformasi untuk menentukan tingkat kerusakannya.",
                "Preparation: Cut away all distorted plate and any buckled girders or stiffeners to sound metal, leaving a margin of good plate (≥¾ thickness).":
                    "Persiapan: Potong semua pelat yang terdistorsi serta gading atau penegar yang menekuk hingga mencapai logam yang sehat, sisakan batas pelat yang baik (≥¾ ketebalan).",
                "Plate replacement: Fabricate a new steel plate (same or greater thickness) shaped to fit the cut-out.":
                    "Penggantian pelat: Fabrikasi pelat baja baru (ketebalan yang sama atau lebih besar) yang dibentuk agar sesuai dengan potongan.",
                "Alignment and welding: Apply continuing weld passes using E7018 electrodes.":
                    "Penyelarasan dan pengelasan: Lakukan lintasan las berlanjut menggunakan elektroda E7018.",
                "Finishing: Inspect the repair by visual and NDT methods.":
                    "Finishing: Inspeksi hasil perbaikan secara visual dan dengan metode NDT.",

                # Hole
                "A hole in the hull plating may result from heavy corrosion perforation or impact damage.":
                    "Lubang pada pelat lambung dapat disebabkan oleh perforasi korosi yang parah atau kerusakan akibat benturan.",
                "Classification practice is to repair holes by replacing the section with new steel.":
                    "Praktik klasifikasi adalah memperbaiki lubang dengan mengganti bagian tersebut dengan baja baru.",
                "Small holes or punctures are typically patched, whereas larger cuts are made into an “insert plate” replacement.":
                    "Lubang kecil or tusukan biasanya ditambal, sedangkan pemotongan yang lebih besar dilakukan untuk penggantian dengan pelat sisipan (insert plate).",
                "Proper edge preparation (bevels and drilled corners) is essential to prevent future cracks.":
                    "Persiapan tepi yang tepat (bevel dan sudut bor) sangat penting untuk mencegah keretakan di masa depan.",
                "Grind or blast the hole edges clean.":
                    "Gerinda atau bersihkan tepi lubang.",
                "Drill a hole at each corner or the end of cracks to stop crack propagation.":
                    "Bor lubang di setiap sudut atau ujung retakan untuk menghentikan penyebaran retakan.",
                "Cut the hole into a regular shape (e.g. rectangle) with straight sides if possible.":
                    "Potong lubang menjadi bentuk yang beraturan (misalnya persegi panjang) dengan sisi-sisi yang lurus jika memungkinkan.",
                "Remove the affected plating back to good steel (at least ¾ thickness).":
                    "Potong pelat yang terkena dampak hingga mencapai baja yang sehat (setidaknya ¾ ketebalan).",
                "Make sure edges of the remaining hull have smooth backing frames or weld seams.":
                    "Pastikan tepi lambung yang tersisa memiliki rangka penyangga atau kampuh las yang halus.",
                "For a small hole, fit a patch plate matching the cut-out and bevel the edges for welding.":
                    "Untuk lubang kecil, pasang pelat tambalan (patch plate) yang sesuai dengan potongan dan buat bevel pada tepiannya untuk pengelasan.",
                "For larger areas (spanning more than one frame), prepare an insert plate.":
                    "Untuk area yang lebih besar (mencakup lebih dari satu gading), siapkan pelat sisipan (insert plate).",
                "The cut lines and welding seams should align with existing weld lines as practical.":
                    "Garis potong dan kampuh las harus sejajar dengan garis las yang ada jika memungkinkan secara praktis.",
                "Tack the patch/insert in place.":
                    "Las ikat tambalan/sisipan pada posisinya.",
                "Weld continuously around its perimeter.":
                    "Las secara kontinu di sekeliling perimeternya.",
                "Patch plates should be fillet-welded on both sides of the hull for strength.":
                    "Pelat tambalan harus dilas sudut (fillet weld) di kedua sisi lambung untuk kekuatan.",
                "Grind welds smooth when done.":
                    "Gerinda hasil las hingga halus setelah selesai.",
                "Inspect welds for completeness (e.g. by dye-penetrant).":
                    "Inspeksi lasan untuk memastikan kelengkapannya (misalnya dengan dye-penetrant).",
                "Prime and paint the new steel patch to protect against corrosion.":
                    "Beri cat dasar (primer) dan cat pada pelat tambalan baja baru untuk melindungi dari korosi.",
                "Round any sharp corners on the repair to prevent stress risers.":
                    "Bulatkan semua sudut tajam pada perbaikan untuk mencegah konsentrasi tegangan.",

                # Hole Colons
                "Preparation: Grind or blast the hole edges clean.":
                    "Persiapan: Gerinda atau bersihkan tepi lubang.",
                "Cut out to sound metal: Remove the affected plating back to good steel (at least ¾ thickness).":
                    "Pemotongan hingga logam sehat: Potong pelat yang terkena dampak hingga mencapai baja yang sehat (setidaknya ¾ ketebalan).",
                "Patch vs insert: For a small hole, fit a patch plate matching the cut-out and bevel the edges for welding.":
                    "Tambalan vs sisipan: Untuk lubang kecil, pasang pelat tambalan (patch plate) yang sesuai dengan potongan dan buat bevel pada tepiannya untuk pengelasan.",
                "Welding: Tack the patch/insert in place.":
                    "Pengelasan: Las ikat tambalan/sisipan pada posisinya.",
                "Finishing: Inspect welds for completeness (e.g. by dye-penetrant).":
                    "Finishing: Inspeksi lasan untuk memastikan kelengkapannya (misalnya dengan dye-penetrant).",

                # Paint Failure
                "Paint failure includes peeling, blistering, or loss of coating integrity on the ship’s surfaces.":
                    "Kegagalan cat meliputi pengelupasan, penggelembungan, atau hilangnya integritas lapisan pelindung pada permukaan kapal.",
                "Causes include UV degradation, moisture ingress under the paint, or abrasion.":
                    "Penyebabnya meliputi degradasi akibat sinar UV, masuknya kelembapan di bawah cat, atau abrasi.",
                "When paint fails, underlying steel may start to corrode.":
                    "Ketika cat gagal, baja di bawahnya dapat mulai korosi.",
                "A typical remediation is to remove the failed coating and apply new protective layers.":
                    "Remediasi umum adalah menghilangkan lapisan cat yang gagal dan mengaplikasikan lapisan pelindung baru.",
                "For example, epoxy or polyurethane paints (with appropriate primers) are commonly used to prevent corrosion and UV damage.":
                    "Sebagai contoh, cat epoxy atau polyurethane (dengan cat dasar yang sesuai) umumnya digunakan untuk mencegah korosi dan kerusakan akibat sinar UV.",
                "Scrape and sand the delaminated paint areas or abrasive-blast to near-white metal.":
                    "Kerok dan ampelas area cat yang terkelupas atau lakukan abrasive blasting hingga logam mendekati putih (near-white metal).",
                "Remove all rust or loose particles so that a clean surface is exposed.":
                    "Bersihkan semua karat atau partikel longgar agar permukaan yang bersih terekspos.",
                "Feather the edges of intact paint to avoid abrupt transitions.":
                    "Buat tepi cat yang masih utuh melandai (feathering) untuk menghindari transisi yang tajam.",
                "If corrosion has started, treat it with rust converter or grind out severely pitted areas.":
                    "Jika korosi telah dimulai, atasi dengan konverter karat (rust converter) atau gerinda area yang berlubang parah.",
                "Wipe down the surface with solvent to remove grease and dust.":
                    "Seka permukaan dengan pelarut (solvent) untuk membersihkan gemuk dan debu.",
                "Apply a coat of zinc-rich or epoxy primer to the prepared steel (this provides corrosion inhibition).":
                    "Aplikasikan satu lapis cat dasar kaya seng (zinc-rich primer) atau primer epoxy pada baja yang telah disiapkan (ini memberikan ketahanan terhadap korosi).",
                "Ensure full coverage of exposed metal.":
                    "Pastikan cakupan penuh pada logam yang terekspos.",
                "Over the primer, apply the specified shipboard paint system (often a high-build epoxy or polyurethane topcoat).":
                    "Di atas cat dasar, aplikasikan sistem cat kapal yang ditentukan (sering kali berupa cat akhir high-build epoxy atau polyurethane).",
                "Use brush, roller or spray to achieve the required thickness.":
                    "Gunakan kuas, kuas rol, atau semprotan untuk mencapai ketebalan yang disyaratkan.",
                "Multiple coats may be needed for full protection.":
                    "Beberapa lapisan mungkin diperlukan untuk perlindungan penuh.",
                "Allow proper drying time for each coat per manufacturer specs.":
                    "Berikan waktu pengeringan yang cukup untuk setiap lapisan sesuai spesifikasi pabrikan.",
                "Inspect the new coating for holidays (pin-holes) and apply touch-ups.":
                    "Inspeksi lapisan baru untuk mendeteksi celah kecil (pin-holes/holidays) dan lakukan perbaikan (touch-up).",
                "Ensure any hull markings or anti-foul topcoats are reapplied if needed.":
                    "Pastikan setiap marka lambung atau lapisan anti-fouling diaplikasikan kembali jika diperlukan.",

                # Paint Failure Colons
                "Removal of failed coating: Scrape and sand the delaminated paint areas or abrasive-blast to near-white metal.":
                    "Pembersihan lapisan cat yang gagal: Kerok dan ampelas area cat yang terkelupas atau lakukan abrasive blasting hingga logam mendekati putih (near-white metal).",
                "Surface preparation: Feather the edges of intact paint to avoid abrupt transitions.":
                    "Persiapan permukaan: Buat tepi cat yang masih utuh melandai (feathering) untuk menghindari transisi yang tajam.",
                "Priming: Apply a coat of zinc-rich or epoxy primer to the prepared steel (this provides corrosion inhibition).":
                    "Pemberian cat dasar (priming): Aplikasikan satu lapis cat dasar kaya seng (zinc-rich primer) atau primer epoxy pada baja yang telah disiapkan (ini memberikan ketahanan terhadap korosi).",
                "Topcoating: Over the primer, apply the specified shipboard paint system (often a high-build epoxy or polyurethane topcoat).":
                    "Pelapisan akhir (topcoating): Di atas cat dasar, aplikasikan sistem cat kapal yang ditentukan (sering kali berupa cat akhir high-build epoxy atau polyurethane).",
                "Curing: Allow proper drying time for each coat per manufacturer specs.":
                    "Pengeringan (curing): Berikan waktu pengeringan yang cukup untuk setiap lapisan sesuai spesifikasi pabrikan.",

                # Fouling
                "Fouling (biofouling) is the accumulation of marine organisms – such as barnacles, algae, and mussels – on the hull surface.":
                    "Fouling (biofouling) adalah akumulasi organisme laut – seperti teritip, alga, dan kerang – pada permukaan lambung.",
                "It increases hydrodynamic drag and fuel consumption.":
                    "Ini meningkatkan hambatan hidrodinamis dan konsumsi bahan bakar.",
                "Ships use anti-fouling coatings (paints with biocides) to minimise this growth, but over time fouling still occurs.":
                    "Kapal menggunakan lapisan anti-fouling (cat dengan biosida) to meminimalkan pertumbuhan ini, namun seiring waktu fouling tetap terjadi.",
                "Causes include biological attachment in calm waters and degraded anti-foul coating.":
                    "Penyebabnya meliputi penempelan biologis di air yang tenang dan degradasi lapisan anti-fouling.",
                "Fouling removal is a regular maintenance task to restore performance.":
                    "Pembersihan fouling adalah tugas perawatan rutin untuk memulihkan kinerja.",
                "In drydock or alongside, inspect hull for fouling.":
                    "Di galangan kapal (drydock) or saat bersandar, inspeksi lambung terhadap fouling.",
                "Determine fouling severity and check coating condition underneath.":
                    "Tentukan tingkat keparahan fouling dan periksa kondisi lapisan cat di bawahnya.",
                "Remove fouling using divers or remotely-operated vehicles equipped with brushes or high-pressure water jets.":
                    "Bersihkan fouling menggunakan penyelam atau kendaraan yang dioperasikan dari jarak jauh (ROV) yang dilengkapi dengan sikat atau jet air bertekanan tinggi.",
                "Earlier methods used divers with brush machines, but newer ROVs can jet-blast the hull efficiently.":
                    "Metode sebelumnya menggunakan penyelam dengan mesin sikat, namun ROV yang lebih baru dapat membersihkan lambung dengan semprotan air (jet-blast) secara efisien.",
                "Ensure the anti-fouling paint is not overly damaged during cleaning.":
                    "Pastikan cat anti-fouling tidak rusak secara berlebihan selama pembersihan.",
                "After scrubbing, power-wash the hull to rinse off residual debris.":
                    "Setelah digosok, cuci lambung dengan air bertekanan tinggi untuk membilas sisa-sia kotoran.",
                "If the anti-fouling coating is worn or abraded, apply a fresh coat.":
                    "Jika lapisan anti-fouling aus atau terkikis, aplikasikan lapisan baru.",
                "Typically, first apply a suitable primer if required, then spray or brush on anti-fouling paint formulated for marine use.":
                    "Biasanya, pertama aplikasikan primer yang sesuai jika diperlukan, lalu semprot atau kuas cat anti-fouling yang diformulasikan untuk penggunaan kelautan.",
                "Allow the paint to cure as per specification.":
                    "Berikan waktu pengeringan bagi cat sesuai spesifikasi.",
                "Return vessel to service promptly since clean hull greatly reduces fuel use.":
                    "Segera kembalikan kapal ke layanan karena lambung yang bersih sangat mengurangi konsumsi bahan bakar.",

                # Fouling Colons
                "Inspection: In drydock or alongside, inspect hull for fouling.":
                    "Inspeksi: Di galangan kapal (drydock) atau saat bersandar, inspeksi lambung terhadap fouling.",
                "Hull cleaning: Remove fouling using divers or remotely-operated vehicles equipped with brushes or high-pressure water jets.":
                    "Pembersihan lambung: Bersihkan fouling menggunakan penyelam atau kendaraan yang dioperasikan dari jarak jauh (ROV) yang dilengkapi dengan sikat atau jet air bertekanan tinggi.",
                "Pressure washing: After scrubbing, power-wash the hull to rinse off residual debris.":
                    "Pencucian bertekanan: Setelah digosok, cuci lambung dengan air bertekanan tinggi untuk membilas sisa-sia kotoran.",
                "Paint renewal: If the anti-fouling coating is worn or abraded, apply a fresh coat.":
                    "Pembaruan cat: Jika lapisan anti-fouling aus atau terkikis, aplikasikan lapisan baru.",
                "Finishing: Allow the paint to cure as per specification.":
                    "Finishing: Berikan waktu pengeringan bagi cat sesuai spesifikasi.",
            }

            # First, clean markdown from lookup key and target string for matching sentence
            parts = re.split(r'(\.\s+)', temp_text)
            translated_parts = []
            for part in parts:
                if not part.strip() or part == '.' or re.match(r'^\.\s+$', part):
                    translated_parts.append(part)
                    continue

                # Strip citation brackets first to match raw sentence
                cleaned_part = re.sub(r'【[^】]+】', '', part)
                cleaned_part = re.sub(r':?contentReference\[[^\]]+\](?:\{[^\}]+\})?', '', cleaned_part)

                # Strip markdown for translation mapping
                cleaned_part_stripped = re.sub(r'\*\*|##|#', '', cleaned_part)
                cleaned_part_stripped = re.sub(r'^\s*-\s*', '', cleaned_part_stripped)
                cleaned_part_stripped = cleaned_part_stripped.strip().rstrip('.')
                cleaned_part_stripped = " ".join(cleaned_part_stripped.split())

                matched_translation = None
                for key, val in sentence_translations.items():
                    key_clean = re.sub(r'\*\*|##|#', '', key)
                    key_clean = re.sub(r'^\s*-\s*', '', key_clean)
                    key_clean = key_clean.strip().rstrip('.')
                    key_clean = " ".join(key_clean.split())
                    if key_clean.lower() == cleaned_part_stripped.lower():
                        matched_translation = val
                        break

                if matched_translation:
                    cites = re.findall(r'【[^】]+】|:?contentReference\[[^\]]+\](?:\{[^\}]+\})?', part)
                    has_period = part.rstrip().endswith('.')
                    cite_str = "".join(cites)

                    base = matched_translation.rstrip('.')
                    if has_period:
                        translated_parts.append(f"{base}{cite_str}.")
                    else:
                        translated_parts.append(f"{base}{cite_str}")
                else:
                    fallback_part = part
                    for k, v in translations.items():
                        if k in fallback_part:
                            fallback_part = fallback_part.replace(k, v)
                    translated_parts.append(fallback_part)

            translated_text = "".join(translated_parts)
            
            translated_text = translated_text.replace(" at ", " pada ")
            translated_text = translated_text.replace("Item: ", "Item: ")
            translated_text = translated_text.replace("Proof image could not be embedded", "Gambar bukti tidak dapat disematkan")
            translated_text = translated_text.replace("Subject to onboard verification", "Tunduk pada verifikasi di atas kapal")

            # Restore protected glossary terms
            for placeholder, original in placeholders.items():
                translated_text = translated_text.replace(placeholder, original)

        # 2. Markdown Cleanup
        final_text = re.sub(r'\*\*|##|#', '', translated_text)
        final_text = re.sub(r'^\s*-\s*', '', final_text)
        final_text = " ".join(final_text.split())

        if is_paragraph:
            print(f"[REPORT] Translated text: {translated_text}")
            print(f"[REPORT] Markdown removed: {final_text}")
            print(f"[REPORT] Final text written to DOCX: {final_text}")

        return final_text

    def load_json(self, path):
        with open(path, "r", encoding="utf-8") as file:
            return json.load(file)

    def _safe_float(self, value, default=0.0):
        try:
            return float(value)
        except (TypeError, ValueError):
            return default

    def _format_currency(self, value, currency):
        if currency == "IDR":
            currency = "IDR"
        return f"{self._safe_float(value):,.2f} {currency}"

    def _clean_label(self, value, fallback):
        if not value:
            return fallback
        return str(value).replace("_", " ").strip().title()

    def _compact_text(self, value, fallback):
        if not value:
            return fallback
        text = " ".join(str(value).split())
        return text[:180].rstrip() if len(text) > 180 else text

    def _categorize_item(self, item_name):
        normalized = str(item_name or "").lower()
        if any(keyword in normalized for keyword in self.LABOR_KEYWORDS):
            return "labor"
        if any(keyword in normalized for keyword in self.EQUIPMENT_KEYWORDS):
            return "equipment"
        return "material"

    def _normalize_item(self, item, default_currency):
        quantity_per_sqm = self._safe_float(item.get("quantity_per_sqm", 0))
        required_quantity = self._safe_float(
            item.get("required_quantity", item.get("quantity", quantity_per_sqm))
        )
        unit_cost = self._safe_float(item.get("unit_cost", item.get("cost", 0)))
        total_cost = required_quantity * unit_cost
        return {
            "item_name": item.get("item_name") or item.get("name") or "New Item",
            "metrics": item.get("metrics") or item.get("unit") or "pcs",
            "quantity_per_sqm": quantity_per_sqm,
            "required_quantity": required_quantity,
            "unit_cost": unit_cost,
            "currency": item.get("currency") or default_currency,
            "total_cost": total_cost,
        }

    def _normalize_report_payload(self, repair_outputs):
        source = repair_outputs if isinstance(repair_outputs, dict) else {}
        repair_summary = source.get("repair_summary") if isinstance(source.get("repair_summary"), dict) else {}
        defect_repairs = source.get("defect_repairs") if isinstance(source.get("defect_repairs"), dict) else {}
        default_currency = repair_summary.get("currency", "IDR")

        normalized_repairs = {}
        total_estimated_cost = 0.0
        total_material_cost = 0.0
        total_labor_cost = 0.0
        total_equipment_cost = 0.0
        severity_distribution = {"low": 0, "medium": 0, "high": 0}

        for defect_id, defect_data in defect_repairs.items():
            estimation = defect_data.get("repair_estimation", {})
            defect_currency = estimation.get("currency", default_currency)
            raw_items = estimation.get("required_items", [])
            required_items = [
                self._normalize_item(item, defect_currency)
                for item in raw_items
                if isinstance(item, dict)
            ]

            defect_total = 0.0
            defect_material = 0.0
            defect_labor = 0.0
            defect_equipment = 0.0

            for item in required_items:
                category = self._categorize_item(item.get("item_name"))
                line_total = self._safe_float(item.get("total_cost"))
                defect_total += line_total
                if category == "labor":
                    defect_labor += line_total
                elif category == "equipment":
                    defect_equipment += line_total
                else:
                    defect_material += line_total

            if not required_items:
                defect_total = self._safe_float(estimation.get("estimated_total_cost", 0))
                defect_material = self._safe_float(estimation.get("material_cost", 0))
                defect_labor = self._safe_float(estimation.get("labor_cost", 0))
                defect_equipment = self._safe_float(estimation.get("equipment_cost", 0))

            severity = str(defect_data.get("severity", "low")).lower()
            severity_distribution[severity] = severity_distribution.get(severity, 0) + 1

            normalized_repairs[defect_id] = {
                **defect_data,
                "repair_estimation": {
                    **estimation,
                    "currency": defect_currency,
                    "required_items": required_items,
                    "estimated_total_cost": round(defect_total, 2),
                    "material_cost": round(defect_material, 2),
                    "labor_cost": round(defect_labor, 2),
                    "equipment_cost": round(defect_equipment, 2),
                },
            }

            total_estimated_cost += defect_total
            total_material_cost += defect_material
            total_labor_cost += defect_labor
            total_equipment_cost += defect_equipment

        return {
            **source,
            "repair_summary": {
                **repair_summary,
                "total_defects": len(normalized_repairs),
                "total_estimated_cost": round(total_estimated_cost, 2),
                "total_material_cost": round(total_material_cost, 2),
                "total_labor_cost": round(total_labor_cost, 2),
                "total_equipment_cost": round(total_equipment_cost, 2),
                "currency": default_currency,
                "severity_distribution": severity_distribution,
            },
            "defect_repairs": normalized_repairs,
        }

    def _set_cell_shading(self, cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        tc_pr.append(shading)

    def _set_table_borders(self, table, color="9FA8B2"):
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        borders = tbl_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = f"w:{edge}"
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "6")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color)

    def _style_table(self, table):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        self._set_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    def _configure_document(self, document):
        section = document.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

        styles = document.styles
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(9)
        styles["Normal"].paragraph_format.space_after = Pt(3)

    def _add_header(self, document, title_text, reference_code, estimate_date):
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(4.8)
        table.columns[1].width = Inches(2.2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]

        brand = left_cell.paragraphs[0]
        brand.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = brand.add_run(self.translate("Marine Technical Services", self.current_lang))
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(31, 54, 91)

        for line in (
            "Hull inspection and repair estimation support",
            "Prepared from digital survey findings",
            "Contact: operations@marine-technical.local"
        ):
            paragraph = left_cell.add_paragraph(self.translate(line, self.current_lang))
            paragraph.paragraph_format.space_after = Pt(1)

        self._set_cell_shading(right_cell, "E9EFF7")
        title = right_cell.paragraphs[0]
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(self.translate(title_text, self.current_lang))
        title_run.bold = True
        title_run.font.size = Pt(14)

        ref_paragraph = right_cell.add_paragraph()
        ref_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_run = ref_paragraph.add_run(f"{self.translate('Ref', self.current_lang)}: {reference_code}")
        ref_run.bold = True

        date_paragraph = right_cell.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run(f"{self.translate('Date', self.current_lang)}: {estimate_date}")

        self._style_table(table)
        document.add_paragraph()

    def _add_metadata(self, document, vessel_name, imo_number, report_scope, currency):
        table = document.add_table(rows=2, cols=4)
        labels = [
            (self.translate("Vessel Name", self.current_lang), vessel_name or self.translate("To be confirmed", self.current_lang)),
            (self.translate("IMO Number", self.current_lang), imo_number or self.translate("To be confirmed", self.current_lang)),
            (self.translate("Document Type", self.current_lang), self.translate("Repair Cost Estimate", self.current_lang)),
            (self.translate("Scope", self.current_lang), self.translate(report_scope, self.current_lang)),
            (self.translate("Prepared By", self.current_lang), self.translate("Automated Inspection Workflow", self.current_lang)),
            (self.translate("Currency", self.current_lang), currency),
            (self.translate("Basis", self.current_lang), self.translate("Inspection findings and repair rules", self.current_lang)),
            (self.translate("Validity", self.current_lang), self.translate("Subject to onboard verification", self.current_lang)),
        ]

        label_fill = "DCE6F1"
        index = 0
        for row in table.rows:
            for _ in range(2):
                label_cell = row.cells[index]
                value_cell = row.cells[index + 1]
                title, value = labels.pop(0)
                label_cell.text = title
                value_cell.text = value
                self._set_cell_shading(label_cell, label_fill)
                for paragraph in label_cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                index += 2
            index = 0

        self._style_table(table)
        document.add_paragraph()

    def _build_scope_rows(self, defect_repairs):
        rows = []
        for number, (defect_id, repair_data) in enumerate(defect_repairs.items(), start=1):
            metadata = repair_data.get("defect_metadata", {})
            estimation = repair_data.get("repair_estimation", {})
            defect_name = self.translate(self._clean_label(repair_data.get("defect_name"), "General Defect"), self.current_lang)
            severity = self.translate(self._clean_label(repair_data.get("severity"), "Low"), self.current_lang)
            area = self._safe_float(metadata.get("defect_area"))
            location = " / ".join(
                self.translate(part.get("part_name", "").replace("_", " ").title(), self.current_lang)
                for part in metadata.get("overlapping_parts", [])
                if part.get("part_name")
            ) or self.translate("General area", self.current_lang)
            process = self.translate(repair_data.get("repair_process") or repair_data.get("description") or "Repair procedure to be confirmed.", self.current_lang)
            required_items = estimation.get("required_items", [])

            if required_items:
                defect_total = max(self._safe_float(estimation.get("estimated_total_cost")), 0.0)
                defect_material = max(self._safe_float(estimation.get("material_cost")), 0.0)
                defect_service = max(
                    self._safe_float(estimation.get("labor_cost")) +
                    self._safe_float(estimation.get("equipment_cost")),
                    0.0,
                )
                for item_index, item in enumerate(required_items, start=1):
                    item_name = self.translate(self._clean_label(item.get("item_name"), "Repair Item"), self.current_lang)
                    quantity = self._safe_float(item.get("required_quantity"))
                    metrics = item.get("metrics") or "unit"
                    item_total = max(self._safe_float(item.get("total_cost")), 0.0)
                    category = self._categorize_item(item_name)
                    if category == "material":
                        material_share = item_total
                        service_cost = 0.0
                    else:
                        material_share = 0.0
                        service_cost = item_total
                    description = (
                        f"{defect_name} {self.translate('at', self.current_lang)} {location}. {process} "
                        f"{self.translate('Item', self.current_lang)}: {item_name} ({quantity:.2f} {self.translate(metrics, self.current_lang)})."
                    )
                    rows.append({
                        "no": number if item_index == 1 else "",
                        "description": description,
                        "service": service_cost,
                        "material": material_share,
                        "total": item_total if item_total > 0 else service_cost + material_share,
                        "severity": severity,
                        "area": area,
                        "defect_id": defect_id,
                    })
            else:
                base_total = self._safe_float(estimation.get("estimated_total_cost"))
                material_cost = self._safe_float(estimation.get("material_cost"))
                service_cost = max(base_total - material_cost, 0.0)
                rows.append({
                    "no": number,
                    "description": f"{defect_name} {self.translate('at', self.current_lang)} {location}. {process}",
                    "service": service_cost,
                    "material": material_cost,
                    "total": base_total,
                    "severity": severity,
                    "area": area,
                    "defect_id": defect_id,
                })

        return rows

    def _add_summary_band(self, document, repair_summary):
        currency = repair_summary.get("currency", "IDR")
        severity = repair_summary.get("severity_distribution", {})
        severity_text = (
            f"{self.translate('Low', self.current_lang)}: {severity.get('low', 0)} | "
            f"{self.translate('Medium', self.current_lang)}: {severity.get('medium', 0)} | "
            f"{self.translate('High', self.current_lang)}: {severity.get('high', 0)}"
        )

        table = document.add_table(rows=2, cols=4)
        content = [
            (self.translate("Total Defects", self.current_lang), str(repair_summary.get("total_defects", 0))),
            (self.translate("Estimated Total", self.current_lang), self._format_currency(repair_summary.get("total_estimated_cost", 0), currency)),
            (self.translate("Material Cost", self.current_lang), self._format_currency(repair_summary.get("total_material_cost", 0), currency)),
            (self.translate("Labor + Equipment", self.current_lang), self._format_currency(
                self._safe_float(repair_summary.get("total_labor_cost", 0)) +
                self._safe_float(repair_summary.get("total_equipment_cost", 0)),
                currency,
            )),
            (self.translate("Severity Mix", self.current_lang), severity_text),
            (self.translate("Issue Count", self.current_lang), str(repair_summary.get("total_defects", 0))),
            (self.translate("Basis of Estimate", self.current_lang), self.translate("AI review and operator-approved quantities", self.current_lang)),
            (self.translate("Status", self.current_lang), self.translate("Budgetary estimate", self.current_lang)),
        ]

        cursor = 0
        for row in table.rows:
            for pair_index in range(0, 4, 2):
                label_cell = row.cells[pair_index]
                value_cell = row.cells[pair_index + 1]
                label, value = content[cursor]
                label_cell.text = label
                value_cell.text = value
                self._set_cell_shading(label_cell, "D9EAD3")
                for paragraph in label_cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                cursor += 1

        self._style_table(table)
        document.add_paragraph()

    def _add_scope_table(self, document, defect_repairs, currency):
        heading = document.add_paragraph()
        heading_run = heading.add_run(self.translate("WORK SCOPE AND COST BREAKDOWN", self.current_lang))
        heading_run.bold = True
        heading_run.font.size = Pt(11)

        table = document.add_table(rows=1, cols=5)
        headers = table.rows[0].cells
        headers[0].text = self.translate("No.", self.current_lang)
        headers[1].text = self.translate("Work Description", self.current_lang)
        headers[2].text = self.translate("Service", self.current_lang)
        headers[3].text = self.translate("Material", self.current_lang)
        headers[4].text = self.translate("Total", self.current_lang)

        for cell in headers:
            self._set_cell_shading(cell, "B4C6E7")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        grand_service = 0.0
        grand_material = 0.0
        grand_total = 0.0

        for row_data in self._build_scope_rows(defect_repairs):
            row = table.add_row().cells
            row[0].text = str(row_data["no"])
            row[1].text = row_data["description"]
            row[2].text = self._format_currency(row_data["service"], currency)
            row[3].text = self._format_currency(row_data["material"], currency)
            row[4].text = self._format_currency(row_data["total"], currency)
            grand_service += row_data["service"]
            grand_material += row_data["material"]
            grand_total += row_data["total"]

        total_row = table.add_row().cells
        total_row[0].merge(total_row[1])
        total_row[0].text = self.translate("TOTAL ESTIMATED COST", self.current_lang)
        total_row[2].text = self._format_currency(grand_service, currency)
        total_row[3].text = self._format_currency(grand_material, currency)
        total_row[4].text = self._format_currency(grand_total, currency)

        for cell in total_row:
            self._set_cell_shading(cell, "FCE5CD")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        self._style_table(table)
        document.add_paragraph()

    def _add_defect_register(self, document, defect_repairs):
        heading = document.add_paragraph()
        heading_run = heading.add_run(self.translate("DEFECT REGISTER", self.current_lang))
        heading_run.bold = True
        heading_run.font.size = Pt(11)

        table = document.add_table(rows=1, cols=5)
        headers = table.rows[0].cells
        headers[0].text = self.translate("Defect ID", self.current_lang)
        headers[1].text = self.translate("Defect Type", self.current_lang)
        headers[2].text = self.translate("Location", self.current_lang)
        headers[3].text = self.translate("Severity", self.current_lang)
        headers[4].text = self.translate("Area", self.current_lang)

        for cell in headers:
            self._set_cell_shading(cell, "D9D2E9")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for defect_id, repair_data in defect_repairs.items():
            metadata = repair_data.get("defect_metadata", {})
            parts = metadata.get("overlapping_parts", [])
            location = " / ".join(
                self.translate(part.get("part_name", "").replace("_", " ").title(), self.current_lang)
                for part in parts
                if part.get("part_name")
            ) or self.translate("General area", self.current_lang)
            row = table.add_row().cells
            row[0].text = str(defect_id)
            row[1].text = self.translate(self._clean_label(repair_data.get("defect_name"), "General Defect"), self.current_lang)
            row[2].text = location
            row[3].text = self.translate(self._clean_label(repair_data.get("severity"), "Low"), self.current_lang)
            row[4].text = f"{self._safe_float(metadata.get('defect_area')):.2f} {self.translate(metadata.get('area_metrics') or 'sqm', self.current_lang)}"

        self._style_table(table)
        document.add_paragraph()

    def _add_defect_proof_section(self, document, defect_repairs):
        heading = document.add_paragraph()
        heading_run = heading.add_run(self.translate("DEFECT PROOF AND COST JUSTIFICATION", self.current_lang))
        heading_run.bold = True
        heading_run.font.size = Pt(11)

        for defect_id, repair_data in defect_repairs.items():
            metadata = repair_data.get("defect_metadata", {})
            estimation = repair_data.get("repair_estimation", {})
            parts = metadata.get("overlapping_parts", [])
            location = " / ".join(
                self.translate(part.get("part_name", "").replace("_", " ").title(), self.current_lang)
                for part in parts
                if part.get("part_name")
            ) or self.translate("General area", self.current_lang)

            title = document.add_paragraph()
            title_run = title.add_run(
                f"{defect_id} - {self.translate(self._clean_label(repair_data.get('defect_name'), 'General Defect'), self.current_lang)}"
            )
            title_run.bold = True

            details = document.add_paragraph()
            details.add_run(self.translate("Location: ", self.current_lang)).bold = True
            details.add_run(location)
            details.add_run(self.translate("   Severity: ", self.current_lang)).bold = True
            details.add_run(self.translate(self._clean_label(repair_data.get("severity"), "Low"), self.current_lang))
            details.add_run(self.translate("   Approved Total: ", self.current_lang)).bold = True
            details.add_run(
                self._format_currency(
                    estimation.get("estimated_total_cost", 0),
                    estimation.get("currency", "IDR"),
                )
            )

            reason = document.add_paragraph()
            reason.add_run(self.translate("Basis: ", self.current_lang)).bold = True
            reason.add_run(self.translate(
                repair_data.get("repair_process") or repair_data.get("description") or "Line items are based on the approved defect repair scope.",
                self.current_lang
            ))

            image_path = metadata.get("best_frame_path")
            temp_image_path = None
            if image_path:
                try:
                    parsed = urlparse(str(image_path))
                    if parsed.scheme in ("http", "https"):
                        response = requests.get(str(image_path), timeout=15)
                        response.raise_for_status()
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_file:
                            temp_file.write(response.content)
                            temp_image_path = temp_file.name
                        document.add_picture(temp_image_path, width=Inches(5.4))
                    elif os.path.exists(image_path):
                        document.add_picture(image_path, width=Inches(5.4))
                except Exception as exc:
                    fallback = document.add_paragraph()
                    fallback.add_run(self.translate("Image note: ", self.current_lang)).bold = True
                    fallback.add_run(f"{self.translate('Proof image could not be embedded', self.current_lang)} ({exc}).")
                finally:
                    if temp_image_path and os.path.exists(temp_image_path):
                        os.unlink(temp_image_path)

            proof_table = document.add_table(rows=1, cols=4)
            proof_headers = proof_table.rows[0].cells
            proof_headers[0].text = self.translate("Approved Item", self.current_lang)
            proof_headers[1].text = self.translate("Qty", self.current_lang)
            proof_headers[2].text = self.translate("Unit Cost", self.current_lang)
            proof_headers[3].text = self.translate("Line Total", self.current_lang)
            for cell in proof_headers:
                self._set_cell_shading(cell, "F4CCCC")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for item in estimation.get("required_items", []):
                row = proof_table.add_row().cells
                row[0].text = self.translate(self._clean_label(item.get("item_name"), "Repair Item"), self.current_lang)
                row[1].text = f"{self._safe_float(item.get('required_quantity')):.2f} {self.translate(item.get('metrics') or 'pcs', self.current_lang)}"
                row[2].text = self._format_currency(item.get("unit_cost", 0), item.get("currency", "IDR"))
                row[3].text = self._format_currency(item.get("total_cost", 0), item.get("currency", "IDR"))

            self._style_table(proof_table)
            document.add_paragraph()

    def _add_notes(self, document):
        heading = document.add_paragraph()
        run = heading.add_run(self.translate("NOTES", self.current_lang))
        run.bold = True
        run.font.size = Pt(11)

        notes = [
            "This document is a clean estimate template generated from inspection findings.",
            "All quantities and locations should be confirmed onboard before commercial issue.",
            "The format intentionally avoids real client identities, billing data, and company names.",
            "Descriptions are kept brief for operational review and approval workflows.",
        ]

        for note in notes:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(self.translate(note, self.current_lang))

    def _build_document(self, title_text, reference_code, vessel_name, repair_summary, defect_repairs, lang="en"):
        self.current_lang = lang
        document = Document()
        self._configure_document(document)
        estimate_date = datetime.now().strftime("%d-%m-%Y")
        currency = repair_summary.get("currency", "IDR")
        if currency == "IDR":
            currency = "IDR"

        self._add_header(document, title_text, reference_code, estimate_date)
        self._add_metadata(
            document,
            vessel_name=vessel_name,
            imo_number=repair_summary.get("imo_number"),
            report_scope="Hull defect repair estimate",
            currency=currency,
        )
        self._add_summary_band(document, repair_summary)
        self._add_scope_table(document, defect_repairs, currency)
        self._add_defect_register(document, defect_repairs)
        self._add_defect_proof_section(document, defect_repairs)
        self._add_notes(document)
        return document

    def _finalize_output(self, output_docx_path):
        from services.supabase_service import supabase_service

        if supabase_service.is_configured():
            try:
                public_url = supabase_service.upload_file(
                    output_docx_path,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                return {
                    "document_docx_url": public_url,
                    "document_pdf_url": public_url,
                }
            except Exception as exc:
                print(f"Failed to upload report to Supabase: {exc}")

        return {
            "document_docx_url": output_docx_path,
            "document_pdf_url": output_docx_path,
        }

    def create_report(self, repair_estimation_json_path, lang="en"):
        if isinstance(repair_estimation_json_path, str):
            payload = self.load_json(repair_estimation_json_path)
        else:
            payload = repair_estimation_json_path
        repair_outputs = self._normalize_report_payload(payload)
        repair_summary = repair_outputs.get("repair_summary", {})
        defect_repairs = repair_outputs.get("defect_repairs", {})
        
        title_text = "REPAIR ESTIMATE" if lang != "bahasa" else "ESTIMASI PERBAIKAN"
        document = self._build_document(
            title_text=title_text,
            reference_code="MI-SINGLE",
            vessel_name=repair_outputs.get("vessel_name") or "Vessel Under Inspection",
            repair_summary=repair_summary,
            defect_repairs=defect_repairs,
            lang=lang,
        )

        filename = "vessel_inspection_report.docx" if lang != "bahasa" else "vessel_inspection_report_bahasa.docx"
        output_docx_path = os.path.join(self.output_folder, filename)
        document.save(output_docx_path)
        return self._finalize_output(output_docx_path)

    def create_batch_report(self, batch_id: str, repair_json_paths: list[str], vessel_name: str, lang="en"):
        repair_payloads = []
        for path in repair_json_paths:
            if not os.path.exists(path):
                continue
            repair_payloads.append(self.load_json(path))

        return self.create_batch_report_from_payloads(batch_id, repair_payloads, vessel_name, lang=lang)

    def create_batch_report_from_payloads(self, batch_id: str, repair_payloads: list[dict], vessel_name: str, lang="en"):
        aggregated_total_defects = 0
        aggregated_total_cost = 0.0
        aggregated_material_cost = 0.0
        aggregated_labor_cost = 0.0
        aggregated_equipment_cost = 0.0
        severity_distribution = {"low": 0, "medium": 0, "high": 0}
        currency = "IDR"
        all_defect_repairs = {}

        for payload in repair_payloads:
            repair_outputs = self._normalize_report_payload(payload or {})
            repair_summary = repair_outputs.get("repair_summary", {})
            aggregated_total_defects += int(repair_summary.get("total_defects", 0))
            aggregated_total_cost += self._safe_float(repair_summary.get("total_estimated_cost", 0))
            aggregated_material_cost += self._safe_float(repair_summary.get("total_material_cost", 0))
            aggregated_labor_cost += self._safe_float(repair_summary.get("total_labor_cost", 0))
            aggregated_equipment_cost += self._safe_float(repair_summary.get("total_equipment_cost", 0))
            currency = repair_summary.get("currency", currency)

            severities = repair_summary.get("severity_distribution", {})
            for key in severity_distribution:
                severity_distribution[key] += int(severities.get(key, 0))

            for defect_id, defect_data in repair_outputs.get("defect_repairs", {}).items():
                all_defect_repairs[defect_id] = defect_data

        combined_summary = {
            "total_defects": aggregated_total_defects,
            "total_estimated_cost": round(aggregated_total_cost, 2),
            "total_material_cost": round(aggregated_material_cost, 2),
            "total_labor_cost": round(aggregated_labor_cost, 2),
            "total_equipment_cost": round(aggregated_equipment_cost, 2),
            "currency": currency,
            "severity_distribution": severity_distribution,
        }

        title_text = "BATCH REPAIR ESTIMATE" if lang != "bahasa" else "ESTIMASI PERBAIKAN BATCH"
        document = self._build_document(
            title_text=title_text,
            reference_code=f"MI-BATCH-{batch_id[:8].upper()}",
            vessel_name=vessel_name or "Fleet Inspection",
            repair_summary=combined_summary,
            defect_repairs=all_defect_repairs,
            lang=lang,
        )

        output_dir = os.path.join("outputs", "batches", batch_id)
        os.makedirs(output_dir, exist_ok=True)
        filename = "combined_vessel_inspection_report.docx" if lang != "bahasa" else "combined_vessel_inspection_report_bahasa.docx"
        output_docx_path = os.path.join(output_dir, filename)
        document.save(output_docx_path)
        return self._finalize_output(output_docx_path)
